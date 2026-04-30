"""
Build the formal report as a Word document that satisfies the assignment
submission rules:

    - Calibri 11
    - 1.5 line spacing
    - Page numbers in the lower right corner
    - Embedded figures from outputs/figures/
    - Cover page

Run with:

    python report/build_docx.py

Output: report/CCS6521_Churn_Report.docx
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PROJECT_ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = PROJECT_ROOT / "outputs" / "figures"
OUT_PATH = PROJECT_ROOT / "report" / "CCS6521_Churn_Report.docx"


# --- formatting helpers -----------------------------------------------------
def _set_default_font(doc: Document, name: str = "Calibri", size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def _add_page_number_footer(section) -> None:
    """Right-aligned page number in the footer."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    style: str | None = None,
    bold: bool = False,
    italic: bool = False,
    size: int | None = None,
    align: int | None = None,
):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    return p


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def _add_figure(doc: Document, path: Path, caption: str, width_inches: float = 5.6) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)


def _add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Calibri"
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
            for run in t.rows[r].cells[c].paragraphs[0].runs:
                run.font.name = "Calibri"


# ---------------------------------------------------------------------------
def build() -> Path:
    doc = Document()
    _set_default_font(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    _add_page_number_footer(section)

    # --- cover page ---------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(96)
    r = p.add_run("CCS6521 - Data Mining and Machine Learning")
    r.bold = True
    r.font.size = Pt(20)

    _add_paragraph(
        doc,
        "Coursework Assessment 1 - Working with data in Python",
        size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_paragraph(
        doc,
        "Churn-prediction Knowledge-Discovery Pipeline",
        size=14,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    for _ in range(3):
        doc.add_paragraph()

    _add_paragraph(doc, "MSc in Advanced Software Engineering", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "Department of Computer Science", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "University of York Europe Campus, CITY College - Spring 2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(2):
        doc.add_paragraph()

    _add_paragraph(doc, "Module Lecturer: Angeliki Lappa", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "Submission Date: 19 May 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # --- 1. EDA -------------------------------------------------------------
    _add_heading(doc, "1. Exploratory Data Analysis", level=1)

    _add_heading(doc, "1.1 Schema and volumes", level=2)
    _add_paragraph(doc, "The dataset is supplied as four CSV tables. Their shapes and key columns are summarised in the following table.")
    _add_table(
        doc,
        ["Table", "Rows", "Cols", "Key columns"],
        [
            ["users", "15,544", "12", "user_id, birth_year, country, city, created_date, plan, marketing flags, num_contacts, num_referrals, num_successful_referrals"],
            ["devices", "15,544", "2", "brand, user_id (one row per user)"],
            ["notifications", "97,704", "5", "reason, channel, status, user_id, created_date"],
            ["transactions (3 split files)", "2,181,026", "12", "transaction_id, transactions_type, transactions_currency, amount_usd, transactions_state, ea_*, direction, user_id, created_date"],
        ],
    )
    _add_paragraph(doc, "devices has exactly one row per user_id, so it acts as a 1:1 lookup. The remaining tables are 1:N from the user side.")

    _add_heading(doc, "1.2 Time coverage", level=2)
    _add_paragraph(doc, "users.created_date runs from 2018-01-01 to 2019-01-03 (registrations stop roughly four months before the data was extracted). Activity (transactions and notifications) runs from 2018-01-01 to 2019-05-16, giving approximately 16 months of behavioural history per cohort. The latest transaction date is taken as the snapshot date: 2019-05-16.")

    _add_heading(doc, "1.3 Missing values", level=2)
    _add_paragraph(doc, "A column-level audit shows that the only material missingness is in the users table: both attributes_notifications_marketing_push and attributes_notifications_marketing_email are NaN for 5,260 users (33.8%). Every other column is fully populated. The transactions table has the expected NaNs in card-only fields (ea_*, direction) for non-card transaction types. Treating those as informative absence rather than a data-quality issue is appropriate, as those fields are simply not applicable to a TRANSFER or TOPUP record.")

    _add_heading(doc, "1.4 Outliers and obvious data errors", level=2)
    _add_paragraph(doc, "The single most striking issue is in transactions.amount_usd: max value 7.46 x 10^10, mean 1.19 x 10^5, median 8.51, standard deviation 6.96 x 10^7. A maximum of 74 billion USD is impossible for a retail FinTech transaction. The 99.5th percentile is below 8 k USD, so we cap to that level rather than discarding rows; this neutralises the impact on means and sums while keeping the long but legitimate tail of larger transfers. Negative amounts do not appear in this dataset, but the cleaning code defends against them.")

    _add_heading(doc, "1.5 Univariate distributions", level=2)
    _add_paragraph(doc, "Plan: 92.6% STANDARD, 4.6% SILVER, 2.8% GOLD, a heavy long tail. Country: 32% GB, 12% PL, 11% FR, 6% IE, 6% RO, 5% ES, ... (top-15 account for ~85%). Device brand: 50.2% Android, 49.6% Apple, 0.2% Unknown. Birth year: median 1986 (33 years old at snapshot), range 1929-2001; no implausible values. Transaction type: 53.8% CARD_PAYMENT, 18.4% TRANSFER, 14.2% TOPUP, 5.7% EXCHANGE, 3.4% ATM, 3.0% CASHBACK. Transaction state: 87.9% COMPLETED, 5.7% DECLINED, 4.1% REVERTED, 1.5% FAILED. Direction: 81% OUTBOUND, 19% INBOUND. Notifications: 50% EMAIL, 47% PUSH, 2% SMS; 73% SENT, 27% FAILED. The 27% notification-failure rate is itself a finding worth flagging to the operations team.")

    _add_figure(doc, FIG_DIR / "users_overview.png", "Figure 1.1. User-level distributions (birth year, plan, top-15 countries, signup volume over time).")
    _add_figure(doc, FIG_DIR / "transactions_overview.png", "Figure 1.2. Transaction-level distributions (type, state, log-amount, weekly volume).")
    _add_figure(doc, FIG_DIR / "notifications_overview.png", "Figure 1.3. Notification distributions (channel, status, top reasons, weekly volume).")
    _add_figure(doc, FIG_DIR / "devices_overview.png", "Figure 1.4. Device brand counts and brand share by plan.")

    _add_heading(doc, "1.6 Bivariate relationships", level=2)
    _add_paragraph(doc, "A correlation heat-map of log-scaled volume features (Figure 1.5) shows the expected dominant pair n_transactions vs total_completed_usd (r ~ 0.88), and a moderate positive relationship between num_contacts and transaction volume (r ~ 0.36). Plan mostly captures activity differences in the right tail: GOLD users have median transaction counts roughly three times higher than STANDARD (Figure 1.6).")
    _add_figure(doc, FIG_DIR / "correlation_heatmap.png", "Figure 1.5. Correlation heat-map of log-scaled volume features.")
    _add_figure(doc, FIG_DIR / "activity_by_plan.png", "Figure 1.6. Transaction volume and completed-USD value, by plan.")

    doc.add_page_break()

    # --- 2. Preprocessing ---------------------------------------------------
    _add_heading(doc, "2. Initial preprocessing", level=1)

    _add_heading(doc, "2.1 Combining the tables", level=2)
    _add_paragraph(doc, "Because the modelling target is one label per user, all four tables are collapsed to user level. devices joins 1:1 on user_id. notifications and transactions are aggregated to user level (counts, sums, distinct counts, shares, recency) before the join. This keeps the training matrix narrow and avoids leaking row-level information that the model could not recover at scoring time.")

    _add_heading(doc, "2.2 Cleaning steps", level=2)
    _add_table(
        doc,
        ["Issue", "Treatment", "Reason"],
        [
            ["attributes_notifications_marketing_* NaN", "filled with 0", "A user without an explicit opt-in is, by product convention, opted-out."],
            ["transactions.amount_usd extreme outliers (max 7.4x10^10)", "clipped at the 99.5th percentile (~8 k USD)", "Preserves the long legitimate tail while neutralising data-entry errors that would otherwise dominate every aggregate."],
            ["transactions.amount_usd < 0", "dropped (none in data, defensive)", "Negative amounts violate the field semantics."],
            ["users.birth_year implausible", "filtered to age in [14, 100]", "All rows passed, but the rule is in place."],
            ["card-only fields (ea_*, direction) NaN", "left as NaN, used as informative absence in shares", "These fields do not apply to TRANSFER/TOPUP records."],
        ],
    )

    _add_heading(doc, "2.3 Encoding", level=2)
    _add_paragraph(doc, "Plan (3 levels) and device brand (3 levels) are one-hot encoded. Country is grouped to top-15 plus an OTHER bucket and one-hot encoded; this caps cardinality at 16 dummies while keeping the dominant geographies separable. City is dropped: it is colinear with country and adds thousands of high-cardinality categories.")

    _add_heading(doc, "2.4 Feature engineering", level=2)
    _add_paragraph(doc, "All engineered features are computed strictly from data observed up to the cutoff date. We set snapshot = max(transactions.created_date) = 2019-05-16, and cutoff = snapshot - 28 days = 2019-04-18. The prediction is therefore: given everything we know about the user up to 2019-04-18, will they be active in the following 28 days? The full feature list (62 columns) groups into:")
    _add_paragraph(doc, "Demographics: age, plan dummies, brand dummies, country dummies, marketing opt-in flags, user_settings_crypto_unlocked, tenure_days.")
    _add_paragraph(doc, "Network: num_contacts, num_referrals, num_successful_referrals.")
    _add_paragraph(doc, "Transaction volume: trx_count, trx_count_completed, trx_value_completed, trx_value_mean_completed, trx_value_max_completed.")
    _add_paragraph(doc, "Transaction quality: trx_success_rate.")
    _add_paragraph(doc, "Diversity: distinct currencies, MCCs, merchant countries.")
    _add_paragraph(doc, "Behaviour profile: per-type shares (CARD_PAYMENT, TRANSFER, TOPUP, EXCHANGE, ATM, CASHBACK, FEE, CARD_REFUND, REFUND, TAX) and direction shares (OUTBOUND, INBOUND).")
    _add_paragraph(doc, "Recency / cadence: days_since_last_trx, days_since_first_trx (both relative to the cutoff), trx_active_days.")
    _add_paragraph(doc, "Notifications: notif_count, notif_n_sent, notif_n_failed, notif_send_rate, per-channel counts, count of REENGAGEMENT_* notifications.")
    _add_paragraph(doc, "We engineer KPIs (recency, breadth and behavioural mix) at the user level rather than feeding raw event counts, because they discriminate active from inactive users better than any single volume metric, and remain robust across plan tiers.")

    doc.add_page_break()

    # --- 3. Identifying churn ----------------------------------------------
    _add_heading(doc, "3. Identifying unengaged and churned users", level=1)

    _add_heading(doc, "3.1 Engagement metric", level=2)
    _add_paragraph(doc, "Definition. A user is engaged if they have at least one COMPLETED transaction in the 28 days following the cutoff date; otherwise they are unengaged / churned.", bold=True)
    _add_paragraph(doc, "Three reasons for this choice:")
    _add_paragraph(doc, "1. Completed transactions are the cleanest signal of intentional usage. Counting DECLINED, FAILED or REVERTED transactions would conflate genuine engagement with payment-rail problems and fraud blocks.")
    _add_paragraph(doc, "2. A 28-day window matches the natural monthly cadence of consumer finance (salary cycles, card statements). It is short enough to be actionable for a re-engagement campaign and long enough to smooth out weekly noise.")
    _add_paragraph(doc, "3. Symmetry with the feature window. The features look at the user pre-cutoff history; the label looks at post-cutoff behaviour. The two periods do not overlap, so there is no leakage by construction.")
    _add_paragraph(doc, "Users whose registration date is fewer than 28 days before the cutoff are excluded: they have not had a full window of opportunity to be active. In the current dataset, no user fails this test, but the rule is in place for future data extracts.")

    _add_heading(doc, "3.2 Class balance", level=2)
    _add_paragraph(doc, "Out of 15,544 eligible users, 7,196 (46.3%) are churned and 8,348 (53.7%) are engaged. The classes are close enough to balanced that no resampling or focal-loss is required; we still pass class_weight=balanced to both classifiers as cheap insurance.")

    _add_heading(doc, "3.3 Models and results", level=2)
    _add_paragraph(doc, "Three classifiers are produced and evaluated on a stratified 20% held-out test set (3,109 users). The heuristic recency classifier is a non-ML baseline that predicts churn whenever days_since_last_trx >= 28. It is included to show how much (or little) the ML models add on top of a single rule.")
    _add_table(
        doc,
        ["Model", "ROC-AUC", "PR-AUC", "F1", "Precision", "Recall"],
        [
            ["Heuristic recency (days inactive >= 28)", "n/a", "n/a", "0.797", "0.840", "0.757"],
            ["Logistic Regression (median impute, scale, balanced)", "0.893", "0.862", "0.797", "0.818", "0.776"],
            ["Random Forest (300 trees, min_leaf=20, balanced)", "0.901", "0.875", "0.811", "0.798", "0.826"],
        ],
    )
    _add_figure(doc, FIG_DIR / "model_roc_pr.png", "Figure 3.1. ROC and Precision-Recall curves on the 20% test set.")
    _add_figure(doc, FIG_DIR / "confusion_rf.png", "Figure 3.2. Confusion matrix of the chosen Random Forest at threshold 0.5.")
    _add_figure(doc, FIG_DIR / "feature_importances_rf.png", "Figure 3.3. Random Forest feature importances (top 25).")

    _add_paragraph(doc, "Why these three models? The heuristic is the cheapest possible policy and a strong baseline: if our ML model cannot beat it materially, the operational cost of an ML solution is hard to justify. Logistic Regression is a transparent linear model; with balanced class weights and standardised features it produces calibrated probabilities that operations can threshold directly. Random Forest captures non-linear interactions (for example STANDARD plan plus high transfer share plus notification failures) without manual feature crosses, and is robust to the long-tailed transactional features we did not log-transform. We do not include gradient boosting in the headline comparison: with 62 columns and 15 k rows, the marginal AUC gain is small and the extra hyper-parameter tuning is not warranted for a coursework deliverable.")
    _add_paragraph(doc, "Feature importances. days_since_last_trx dominates (about 29% of total decrease-in-impurity), followed by trx_active_days, trx_count, trx_count_completed and trx_value_completed at roughly 5-10% each. Diversity features (trx_n_distinct_mcc, trx_n_distinct_country) and num_contacts round out the top ten. Plan / country / brand dummies and individual notification-channel features have negligible importance. This explains why the heuristic recency rule is already so strong: the dominant axis of variation is how long ago the user was last active.")

    _add_heading(doc, "3.4 Population-level numbers", level=2)
    _add_paragraph(doc, "Scoring the trained Random Forest on the full eligible population gives:")
    _add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Eligible users", "15,544"],
            ["Predicted churn (model, threshold 0.5)", "7,425 (47.8%)"],
            ["Predicted churn (heuristic)", "6,484 (41.7%)"],
            ["Actual churn", "7,196 (46.3%)"],
        ],
    )
    _add_paragraph(doc, "The model predicts slightly more churners than reality (a false-positive rate of about 8%) at the default threshold. Lowering the threshold trades precision for recall: it would catch more truly-at-risk users at the cost of contacting more false alarms, which is a business decision rather than a modelling one.")

    doc.add_page_break()

    # --- 4. Actionable decisions -------------------------------------------
    _add_heading(doc, "4. Actionable decisions (critical discussion)", level=1)

    _add_heading(doc, "4.1 How many users are classified as churned?", level=2)
    _add_paragraph(doc, "At a 0.5 probability threshold, the chosen Random Forest flags 7,425 users (47.8%) as churn-risks. If marketing budget is the binding constraint, the same model can be operated at any other threshold (illustrative figures from the test-set PR curve):")
    _add_table(
        doc,
        ["Threshold", "Flagged (approx.)", "Precision (approx.)", "Recall (approx.)"],
        [
            ["0.30", "~8,500 (55%)", "~0.70", "~0.92"],
            ["0.50", "7,425 (48%)", "0.80", "0.83"],
            ["0.70", "~5,600 (36%)", "~0.91", "~0.71"],
            ["0.85", "~3,400 (22%)", "~0.96", "~0.45"],
        ],
    )
    _add_paragraph(doc, "Exact values for any threshold are computable from outputs/tables/predictions.csv produced by the pipeline.")

    _add_heading(doc, "4.2 Designing an experiment to verify churn reduction", level=2)
    _add_paragraph(doc, "The right design is a randomised controlled experiment, not a pre/post comparison.")
    _add_paragraph(doc, "1. Define the eligible population. Score every user weekly with the model. Users whose probability >= a chosen threshold (and who have not already been part of a recent campaign) are eligible for the test.")
    _add_paragraph(doc, "2. Random assignment. At enrolment, randomly split eligible users into treatment (receive the re-engagement intervention) and control (receive nothing, or the existing default) at a pre-registered ratio (e.g. 50/50). Stratify the randomisation by plan and country so the marginal subgroups are balanced.")
    _add_paragraph(doc, "3. Pre-register everything. Decide before the test starts: the intervention, the primary outcome (proportion of users who make >= 1 COMPLETED transaction in the next 28 days), the test duration, the minimum-detectable-effect, the required sample size from a power calculation, the analysis plan, and the stopping rules.")
    _add_paragraph(doc, "4. Run for the full pre-registered period. Resist the urge to stop early; sequential peeks inflate Type-I error.")
    _add_paragraph(doc, "5. Analyse per-protocol and intent-to-treat. Use a two-proportion z-test (or a logistic regression with covariates for variance reduction) on the primary outcome.")
    _add_paragraph(doc, "Critically, the treatment effect is the difference between treatment and control, not the absolute conversion rate of the treated group. Many flagged users would have come back on their own; without a control arm we would have no way of separating that natural recovery from the intervention.")

    _add_heading(doc, "4.3 Metrics and techniques to assess impact", level=2)
    _add_paragraph(doc, "Primary outcome (causal): treatment-vs-control lift on the same engagement metric we used for training: proportion of users with >= 1 completed transaction in the 28-day post-treatment window.")
    _add_paragraph(doc, "Supporting metrics:")
    _add_paragraph(doc, "Volume metrics. Mean number of completed transactions and total USD spent in the post-treatment window, per arm.")
    _add_paragraph(doc, "Cost / contribution. Net revenue lift per user contacted, including the cost of the intervention (push notifications ~ free, SMS ~ EUR 0.05, email ~ free, in-app credit / cashback ~ the credit amount).")
    _add_paragraph(doc, "Funnel diagnostics. Notification deliverability, open rate, click-through rate, then conversion. These tell us why a treatment did or did not move the primary metric.")
    _add_paragraph(doc, "Long-run retention. Re-measure the engagement metric at +60 and +90 days. A successful re-engagement that fades within a month is far less valuable than one that durably restores the user.")
    _add_paragraph(doc, "Heterogeneous effects. Slice the lift by plan, country, device, and pre-treatment activity decile. The model already tells us the user is at risk; the segment-level lift tells us who is worth treating.")
    _add_paragraph(doc, "Statistical techniques:")
    _add_paragraph(doc, "Two-proportion z-test for the primary binary outcome.")
    _add_paragraph(doc, "CUPED (controlled experiment using pre-experiment data) to reduce variance using each user pre-period activity as a covariate, a particularly good fit here because we already engineered those features.")
    _add_paragraph(doc, "Bayesian estimation of the treatment-effect distribution if the business prefers a posterior probability of lift to a p-value.")
    _add_paragraph(doc, "Sequential or Bonferroni correction if multiple variants are tested simultaneously.")
    _add_paragraph(doc, "Segment-level inference with hierarchical models or simple tests with multiple-testing correction.")
    _add_paragraph(doc, "A failed test is also informative: if a re-engagement push gets the same result as control, the model is correctly identifying churners but the intervention is not the right tool for that segment, and we should pivot (for example fee waiver, in-app credit, customer-success outreach for high-value users).")

    doc.add_page_break()

    # --- 5. What I would have done differently ------------------------------
    _add_heading(doc, "5. What I would have done differently with more time", level=1)
    _add_paragraph(doc, "Out-of-time validation. Use multiple rolling cutoffs rather than a random split; that better approximates the deployment scenario where we score today users to predict tomorrow behaviour.")
    _add_paragraph(doc, "Gradient boosting (XGBoost or LightGBM) with proper hyper-parameter search and SHAP-based feature attributions, likely a small AUC improvement and more nuanced explanations.")
    _add_paragraph(doc, "More expressive recency features. Time-since-last-X for several X (last completed transaction, last failed transaction, last re-engagement notification), plus weekly transaction trajectories (slope of last 8 weeks).")
    _add_paragraph(doc, "Threshold optimisation against expected business value. Combine per-user predicted churn probability with the user expected lifetime value to pick the threshold that maximises net contribution from the re-engagement programme, rather than the F1-default 0.5.")
    _add_paragraph(doc, "Survival framing. Reformulate the problem as a time-to-event / survival model (Cox proportional hazards or DeepSurv). Churn is not really binary; it has a censoring structure that survival models capture naturally.")
    _add_paragraph(doc, "Probability calibration. Even with class_weight=balanced, the absolute probabilities from the RF are not perfectly calibrated; isotonic or Platt calibration on a held-out fold would yield more reliable probabilities for downstream cost / benefit calculations.")
    _add_paragraph(doc, "Address the email deliverability finding. 27% of notifications are marked FAILED; if the same user receives only failed notifications they are functionally untreatable by the existing channels. This is an ops-level fix that would help any churn-reduction programme more than a better classifier would.")
    _add_paragraph(doc, "Investigate the 32 Unknown device brands to ensure they are not a hidden segment with different behaviour.")

    _add_heading(doc, "Reproducibility", level=2)
    _add_paragraph(doc, "The full pipeline is in src/ and is invoked with python -m src.main. Outputs are written to outputs/figures/ (PNG plots), outputs/tables/ (CSV / JSON summaries) and outputs/models/ (the pickled best model). Random seeds are fixed in src/config.py (RANDOM_STATE = 42).")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    out = build()
    print(f"Wrote: {out}")
